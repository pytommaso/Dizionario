#### about CSV
import csv

f = open('/Users/PC/Desktop/dictionary/00_word_BOX.csv', newline='')
reader = csv.reader(f)

#### about English dictionary
import enchant
eng = enchant.Dict("en_US")


#### Cleaning list
words = []
for row in reader:
    word = (''.join(row))  #trasformo in stringa effettiva
    if word == '':
        continue
    words.append(word)
cleaned_list = list(set(words))


#### the LOGIC
eng_words = []
other_words = []
locutions = []
for item in cleaned_list:
    if ' ' in item:
        locutions.append(item)
    elif eng.check(item) == True:
        eng_words.append(item.capitalize())
    else:
        other_words.append(item)

eng_words_clean = list(set(eng_words))
eng_words_clean.sort() #questo è un esempio di funzione che va a modificare un parametro esterno....
                    #dopo il suo "passaggio" il parametro sarà diverso da prima

###CREA uno stringone con item di eng_words a capo
b =''
for item in eng_words_clean:
    b = b + item + '\n'

###TRADUZIONE
from googletrans import Translator  # Import Translator module
translator = Translator() # Create object of Translator
translated = translator.translate(b, src='en', dest = 'it') # Translate

##STRUTTURAZIONE DEI DATI
translated_string = translated.text.lower()  # stringa di Tradotti
ital_words = translated_string.split('\n')  # lista di Tradotti

if len(ital_words) == len(eng_words_clean):
    print ('OK, STESSA QUANTITA\'!     parole tradotte/parole inglesi')

records = []
x = 0
for element in eng_words_clean:
    a = []
    a.append(element)
    a.append(ital_words[x])
    records.append(a)
    x += 1

#####   DOCX part
from docx import Document
from docx.shared import Inches, Pt
document = Document()
#style
style = document.styles['Normal']
font = style.font
font.name = 'times new roman'
font.size = Pt(22)
document.add_heading('', 0)

table = document.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells

for x, y in records:
    row_cells = table.add_row().cells
    row_cells[0].text = x
    row_cells[1].text = '\u2192'
    row_cells[2].text = y

document.add_page_break()
document.save('/Users/PC/Desktop/ENG-ITA.docx')

#####   DOCX 2
document2 = Document()
#style
style = document2.styles['Normal']
font = style.font
font.name = 'times new roman'
font.size = Pt(23)
document2.add_heading('', 0)

table = document2.add_table(rows=1, cols=1)

hdr_cells = table.rows[0].cells

for element in locutions:
    row_cells = table.add_row().cells
    row_cells[0].text = element

row_cells = table.add_row().cells

for element in other_words:
    row_cells = table.add_row().cells
    row_cells[0].text = element

document2.add_page_break()
document2.save('/Users/PC/Desktop/Locuzioni Inglesi.docx')
