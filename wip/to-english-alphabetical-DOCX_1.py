#### about CSV
import csv
f = open('list_of_words2.csv', newline='')
reader = csv.reader(f)

#### about English dictionary
import enchant
eng = enchant.Dict("en_US")

#### the LOGIC
eng_words = []
other_words = []
locutions = []
for row in reader:
    word = (''.join(row))  #trasformo in stringa effettiva
    if word == '':
        continue
    if ' ' in word:
        locutions.append(word)
    elif eng.check(word) == True:
        eng_words.append(word.capitalize())
    else:
        other_words.append(word)

eng_words_clean = list(set(eng_words))
eng_words_clean.sort() #questo è un esempio di funzione che va a modificare un parametro esterno....
                    #dopo il suo "passaggio" il parametro sarà diverso da prima

###CREA uno stringone con item di eng_words a capo
b =''
for item in eng_words_clean:
    b = b + item + '\n'

###TRADUZIONE
from googletrans import Translator  # Import Translator module
translator = Translator() # Create object of Translator
translated = translator.translate(b, src='en', dest = 'it') # Translate

##STRUTTURAZIONE DEI DATI
translated_string = translated.text.lower()  # stringa di Tradotti
ital_words = translated_string.split('\n')  # lista di Tradotti

if len(ital_words) == len(eng_words_clean):
    print ('OK, STESSA QUANTITA\'!     parole tradotte/parole inglesi')

records = []
x = 0
for element in eng_words_clean:
    a = []
    a.append(element)
    a.append(ital_words[x])
    records.append(a)
    x += 1

#####   DOCX part
from docx import Document
from docx.shared import Inches, Pt
document = Document()
#style
style = document.styles['Normal']
font = style.font
font.name = 'times new roman'
font.size = Pt(27)
document.add_heading('', 0)

table = document.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells

for x, y in records:
    row_cells = table.add_row().cells
    row_cells[0].text = x
    row_cells[1].text = '\u2192'
    row_cells[2].text = y

#######################      table(autofit = True)
document.add_page_break()
document.save('English_Glossary.docx')
